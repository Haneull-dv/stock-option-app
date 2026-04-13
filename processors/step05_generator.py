"""
Step05 예탁원 신주발행의뢰 문서 자동 생성.

발행가액별로 폴더 구조 생성:
- (전자등록)_<가액>원.hwpx
- 발행등록확인신청서.pdf
- 법인인감증명서.pdf
- <가액>원/
    - (붙임1) 일괄발행등록_세부내역.xlsx
    - (붙임2) 일괄발행등록_의뢰확약서.pdf
    - (붙임3) 발행근거_정관.pdf
    - (붙임4) 주주총회의사록 <가액>원/ (폴더)
    - (붙임5) 주식매수선택권_부여및행사내역서.pdf
    - (붙임6) 배정자_실명확인증표.pdf
    - (붙임7) 배정자_증권계좌사본.pdf
    - (붙임8) 주식납입금보관증명서.pdf (또는 .jpg)
    - (붙임9) 법인등기부등본.pdf

최종 출력: 발행가액별 ZIP 파일
"""
import os
import shutil
from datetime import datetime
from processors.hwpx_writer import generate_hwpx
from processors.pdf_merger import merge_pdfs_in_order
from processors.excel_writer import generate_issuance_detail_excel
from processors.docx_to_pdf import convert_docx_to_pdf, convert_image_to_pdf
from processors.zip_utils import create_zip_from_folder, copy_folder_contents

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATES_DIR = os.path.join(BASE_DIR, 'templates_step05')


def generate_step05_zip(round_obj, applicants, price, config, attachment8_file, output_base_dir):
    """
    발행가액별 신주발행의뢰 폴더 생성 + ZIP.

    Returns:
        dict: {
            'success': bool,
            'zip_path': str,
            'files': [list of generated files],
            'message': str
        }
    """
    price_str = f"{price}"
    price_display = f"{price:,}원"

    # 임시 작업 폴더
    work_dir = os.path.join(output_base_dir, f'_temp_{price}')
    price_folder = os.path.join(work_dir, f'{price}원')
    os.makedirs(price_folder, exist_ok=True)

    files_created = []
    errors = []

    try:
        # ────────────────────────────────────────────────────────
        # 1. (전자등록) HWPX 생성
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [1/12] (전자등록) HWPX 생성 중...")
            electronic_hwpx = _generate_electronic_registration(
                work_dir, price, applicants, config, round_obj
            )
            files_created.append(('전자등록', electronic_hwpx))
            print(f"      ✓ 완료: {os.path.basename(electronic_hwpx)}")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"전자등록: {e}")

        # ────────────────────────────────────────────────────────
        # 2. 발행등록확인신청서 PDF 생성
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [2/12] 발행등록확인신청서 PDF 생성 중...")
            confirmation_pdf = _generate_registration_confirmation(
                work_dir, price, applicants, config, round_obj
            )
            files_created.append(('발행등록확인신청서', confirmation_pdf))
            print(f"      ✓ 완료")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"발행등록확인신청서: {e}")

        # ────────────────────────────────────────────────────────
        # 3. 법인인감증명서 복사
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [3/12] 법인인감증명서 복사 중...")
            corp_seal_pdf = _copy_corporate_seal(work_dir)
            files_created.append(('법인인감증명서', corp_seal_pdf))
            print(f"      ✓ 완료")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"법인인감증명서: {e}")

        # ────────────────────────────────────────────────────────
        # 4. 붙임1: 일괄발행등록 세부내역 XLSX
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [4/12] 붙임1: 일괄발행등록 세부내역 XLSX 생성 중...")
            attachment1_xlsx = os.path.join(
                price_folder, f'(붙임1) 일괄발행등록_세부내역_{price}원.xlsx'
            )
            stock_code = config.get('stock_code', '488280')
            generate_issuance_detail_excel(applicants, price, stock_code, attachment1_xlsx)
            files_created.append(('붙임1', attachment1_xlsx))
            print(f"      ✓ 완료 ({len(applicants)}명)")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"붙임1: {e}")

        # ────────────────────────────────────────────────────────
        # 5. 붙임2: 일괄발행등록 의뢰확약서 HWPX
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [5/12] 붙임2: 일괄발행등록 의뢰확약서 생성 중...")
            total_shares = sum(ap.get('quantity', 0) for ap in applicants)
            attachment2_hwpx = _generate_attachment2(price_folder, config, len(applicants), total_shares)
            files_created.append(('붙임2', attachment2_hwpx))
            print(f"      ✓ 완료")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"붙임2: {e}")

        # ────────────────────────────────────────────────────────
        # 6. 붙임3: 발행근거_정관 복사
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [6/12] 붙임3: 발행근거_정관 복사 중...")
            attachment3_pdf = _copy_attachment3(price_folder)
            files_created.append(('붙임3', attachment3_pdf))
            print(f"      ✓ 완료")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"붙임3: {e}")

        # ────────────────────────────────────────────────────────
        # 7. 붙임4: 주주총회의사록 폴더 복사
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [7/12] 붙임4: 주주총회의사록 폴더 복사 중...")
            attachment4_folder = _copy_attachment4(price_folder, price)
            files_created.append(('붙임4', attachment4_folder))
            file_count = len([f for f in os.listdir(attachment4_folder) if os.path.isfile(os.path.join(attachment4_folder, f))])
            print(f"      ✓ 완료 ({file_count}개 파일)")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"붙임4: {e}")

        # ────────────────────────────────────────────────────────
        # 8. 붙임5: 주식매수선택권 부여및행사내역서 복사
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [8/12] 붙임5: 주식매수선택권 부여및행사내역서 복사 중...")
            attachment5_pdf = _copy_attachment5(price_folder, price)
            files_created.append(('붙임5', attachment5_pdf))
            print(f"      ✓ 완료")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"붙임5: {e}")

        # ────────────────────────────────────────────────────────
        # 9. 붙임6: 배정자 실명확인증표 (신분증 합본)
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [9/12] 붙임6: 신분증 합본 PDF 생성 중...")
            attachment6_pdf = _merge_id_copies(price_folder, applicants, price)
            files_created.append(('붙임6', attachment6_pdf))
            print(f"      ✓ 완료 ({len(applicants)}명)")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"붙임6: {e}")

        # ────────────────────────────────────────────────────────
        # 10. 붙임7: 배정자 증권계좌사본 합본
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [10/12] 붙임7: 계좌사본 합본 PDF 생성 중...")
            attachment7_pdf = _merge_account_copies(price_folder, applicants, price)
            files_created.append(('붙임7', attachment7_pdf))
            print(f"      ✓ 완료 ({len(applicants)}명)")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"붙임7: {e}")

        # ────────────────────────────────────────────────────────
        # 11. 붙임8: 주식납입금 보관증명서 (업로드 파일)
        # ────────────────────────────────────────────────────────
        if attachment8_file:
            try:
                print(f"    [11/12] 붙임8: 주식납입금보관증명서 복사 중...")
                attachment8_dest = os.path.join(price_folder, f'(붙임8) 주식납입금보관증명서_{price}원.pdf')
                ext = os.path.splitext(attachment8_file)[1].lower()
                if ext in ['.jpg', '.jpeg', '.png']:
                    # 이미지 → PDF 변환
                    convert_image_to_pdf(attachment8_file, attachment8_dest)
                    print(f"      ✓ 완료 (이미지→PDF 변환)")
                else:
                    # PDF 그대로 복사
                    shutil.copy2(attachment8_file, attachment8_dest)
                    print(f"      ✓ 완료")
                files_created.append(('붙임8', attachment8_dest))
            except Exception as e:
                print(f"      ✗ 실패: {e}")
                errors.append(f"붙임8: {e}")
        else:
            print(f"    [11/12] 붙임8: 업로드 안됨, 건너뜀")

        # ────────────────────────────────────────────────────────
        # 12. 붙임9: 법인등기부등본 복사
        # ────────────────────────────────────────────────────────
        try:
            print(f"    [12/12] 붙임9: 법인등기부등본 복사 중...")
            attachment9_pdf = _copy_attachment9(price_folder)
            files_created.append(('붙임9', attachment9_pdf))
            print(f"      ✓ 완료")
        except Exception as e:
            print(f"      ✗ 실패: {e}")
            errors.append(f"붙임9: {e}")

        # ────────────────────────────────────────────────────────
        # ZIP 생성
        # ────────────────────────────────────────────────────────
        print(f"    [ZIP] 압축 파일 생성 중...")
        zip_name = f'신주발행의뢰_{price}원.zip'
        zip_path = os.path.join(output_base_dir, zip_name)
        create_zip_from_folder(work_dir, zip_path)
        print(f"      ✓ 완료: {zip_name}")

        # 임시 폴더 삭제
        shutil.rmtree(work_dir, ignore_errors=True)

        return {
            'success': True,
            'zip_path': zip_path,
            'zip_name': zip_name,
            'files': files_created,
            'errors': errors,
            'message': f'{price_display} 폴더 생성 완료'
        }

    except Exception as e:
        # 임시 폴더 삭제
        shutil.rmtree(work_dir, ignore_errors=True)
        return {
            'success': False,
            'message': str(e),
            'errors': errors
        }


# ═══════════════════════════════════════════════════════════════════════════════
# 개별 문서 생성 함수들
# ═══════════════════════════════════════════════════════════════════════════════

def _generate_electronic_registration(work_dir, price, applicants, config, round_obj):
    """전자등록 HWPX 생성."""
    template = os.path.join(TEMPLATES_DIR, '전자등록.hwpx')
    output = os.path.join(work_dir, f'(전자등록)_{price}원.hwpx')

    # 치환할 데이터
    payment_date = config.get('payment_date', '').replace('-', '.')
    dividend_date = config.get('dividend_base_date', '').replace('-', '.')
    listing_date = config.get('listing_date', '').replace('-', '.')

    total_shares = sum(ap.get('quantity', 0) for ap in applicants)

    replacements = {
        '{{납입일}}': payment_date or '',
        '{{배당기산일}}': dividend_date or '',
        '{{유통(상장)예정일}}': listing_date or '',
        '{{행사주식수}}': f'{total_shares:,}',
        '{{행사가액}}': f'{price:,}',
    }

    print(f"        · 납입일: {payment_date or '(빈칸)'}")
    print(f"        · 배당기산일: {dividend_date or '(빈칸)'}")
    print(f"        · 유통(상장)예정일: {listing_date or '(빈칸)'}")
    print(f"        · 행사주식수: {total_shares:,}주")
    print(f"        · 행사가액: {price:,}원")

    generate_hwpx(template, output, replacements)
    return output


def _replace_paragraph_text(paragraph, new_text):
    """
    paragraph의 텍스트를 new_text로 교체하되 첫 번째 run의 서식을 유지.
    """
    # 첫 번째 run의 스타일 저장
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font = first_run.font
        bold = font.bold
        italic = font.italic
        size = font.size
        name = font.name
    else:
        bold = italic = size = name = None

    # 모든 runs를 완전히 제거
    for _ in range(len(paragraph.runs)):
        paragraph._element.remove(paragraph.runs[0]._element)

    # 새 run 추가
    new_run = paragraph.add_run(new_text)

    # 스타일 복원
    if bold is not None:
        new_run.font.bold = bold
    if italic is not None:
        new_run.font.italic = italic
    if size is not None:
        new_run.font.size = size
    if name is not None:
        new_run.font.name = name


def _generate_registration_confirmation(work_dir, price, applicants, config, round_obj):
    """발행등록확인신청서 DOCX → PDF 생성."""
    from docx import Document
    import re

    template = os.path.join(TEMPLATES_DIR, '20260223 발행등록확인신청서_스톡옵션행사_에스투더블유.docx')
    temp_docx = os.path.join(work_dir, f'_temp_발행등록확인신청서_{price}.docx')
    output_pdf = os.path.join(work_dir, '발행등록확인신청서.pdf')

    # 기준일, 발행주식수, 유통예정일 데이터 준비
    payment_date_raw = config.get('payment_date', '')
    listing_date_raw = config.get('listing_date', '')
    total_shares = sum(ap.get('quantity', 0) for ap in applicants)

    # 발행가액별 회차 번호 매핑 (고정)
    price_to_round_number = {
        1250: 15,
        2000: 16,
        4130: 17
    }
    round_number = price_to_round_number.get(price, '')

    # 날짜 형식 변환: 2026-02-23 → 2026년 02월 23일
    def format_korean_date(date_str):
        if not date_str:
            return ''
        parts = date_str.split('-')
        if len(parts) == 3:
            return f'{parts[0]}년 {parts[1]}월 {parts[2]}일'
        return date_str

    payment_date_kr = format_korean_date(payment_date_raw)
    listing_date_kr = format_korean_date(listing_date_raw)

    # DOCX 열기
    doc = Document(template)

    # 모든 단락 텍스트 치환 (runs 기반으로 서식 유지)
    for paragraph in doc.paragraphs:
        full_text = paragraph.text

        # "2. 기 준 일 : 2026년 02월 23일" 형태 치환
        if '기 준 일' in full_text or '기준일' in full_text:
            new_text = re.sub(r'(\d{4}년\s*\d{2}월\s*\d{2}일)', payment_date_kr or '', full_text)
            _replace_paragraph_text(paragraph, new_text)

        # "3. 발행횟수 : 제 회" 형태 치환
        if '발행횟수' in full_text and round_number:
            new_text = re.sub(r'제\s*\d*\s*회', f'제 {round_number}회', full_text)
            _replace_paragraph_text(paragraph, new_text)

        # "4. 발행주식수 : 10,500주" 형태 치환
        if '발행주식수' in full_text:
            new_text = re.sub(r'발행주식수\s*:\s*[\d,]+주', f'발행주식수 : {total_shares:,}주', full_text)
            _replace_paragraph_text(paragraph, new_text)

        # "5. 유통(상장)예정일 : 2026년 월 일" 형태 치환
        if '유통' in full_text and '예정일' in full_text:
            new_text = re.sub(r'(\d{4}년\s*\d*월\s*\d*일)', listing_date_kr or '', full_text)
            _replace_paragraph_text(paragraph, new_text)

    doc.save(temp_docx)

    print(f"        · 기준일: {payment_date_kr or '(빈칸)'}")
    print(f"        · 발행횟수: 제 {round_number}회" if round_number else "        · 발행횟수: (매핑 없음)")
    print(f"        · 발행주식수: {total_shares:,}주")
    print(f"        · 유통예정일: {listing_date_kr or '(빈칸)'}")

    convert_docx_to_pdf(temp_docx, output_pdf)
    os.remove(temp_docx)  # 임시 docx 삭제
    return output_pdf


def _copy_corporate_seal(work_dir):
    """법인인감증명서 복사."""
    src = os.path.join(TEMPLATES_DIR, 'reference_docs', '법인인감증명서.pdf')
    dest = os.path.join(work_dir, '법인인감증명서.pdf')
    shutil.copy2(src, dest)
    return dest


def _generate_attachment2(price_folder, config, applicants_count, total_shares):
    """붙임2: 일괄발행등록 의뢰확약서 HWPX."""
    template = os.path.join(TEMPLATES_DIR, '붙임 2.hwpx')
    temp_hwpx = os.path.join(price_folder, '_temp_붙임2.hwpx')
    output_hwpx = os.path.join(price_folder, '(붙임2) 일괄발행등록_의뢰확약서.hwpx')

    agent_name = config.get('agent_name', '')
    agent_phone = config.get('agent_phone', '')
    agent_rrn = config.get('agent_rrn', '')
    agent_address = config.get('agent_address', '')

    replacements = {
        '{{주주수}}': str(applicants_count),
        '{{주식수}}': f'{total_shares:,}',
        '{{대리인성명}}': agent_name,
        '{{대리인연락처}}': agent_phone,
        '{{대리인주민번호}}': agent_rrn,
        '{{대리인주소}}': agent_address,
    }

    print(f"        · 주주수: {applicants_count}명 (자동계산)")
    print(f"        · 주식수: {total_shares:,}주 (자동계산)")
    print(f"        · 대리인: {agent_name or '(빈칸)'}")

    generate_hwpx(template, temp_hwpx, replacements)

    shutil.move(temp_hwpx, output_hwpx)
    return output_hwpx


def _copy_attachment3(price_folder):
    """붙임3: 발행근거_정관 복사."""
    src = os.path.join(TEMPLATES_DIR, 'reference_docs', '(붙임3) 발행근거_정관_제10조제2항제2호_에스투더블유.pdf')
    dest = os.path.join(price_folder, '(붙임3) 발행근거_정관.pdf')
    shutil.copy2(src, dest)
    return dest


def _copy_attachment4(price_folder, price):
    """붙임4: 주주총회의사록 폴더 복사."""
    src_folder = os.path.join(TEMPLATES_DIR, 'price_specific', str(price), f'(붙임4) 주주총회의사록 {price}원')
    dest_folder = os.path.join(price_folder, f'(붙임4) 주주총회의사록 {price}원')

    if os.path.isdir(src_folder):
        copy_folder_contents(src_folder, dest_folder)
    else:
        raise FileNotFoundError(f"붙임4 폴더를 찾을 수 없습니다: {src_folder}")

    return dest_folder


def _copy_attachment5(price_folder, price):
    """붙임5: 주식매수선택권 부여및행사내역서 복사."""
    src = os.path.join(TEMPLATES_DIR, 'price_specific', str(price), f'(붙임5) 주식매수선택권_부여및행사내역서_에스투더블유 {price}.pdf')
    if not os.path.isfile(src):
        src = os.path.join(TEMPLATES_DIR, 'price_specific', str(price), f'(붙임5) 주식매수선택권_부여및행사내역서_에스투더블유 {price}원.pdf')

    dest = os.path.join(price_folder, f'(붙임5) 주식매수선택권_부여및행사내역서.pdf')

    if os.path.isfile(src):
        shutil.copy2(src, dest)
    else:
        raise FileNotFoundError(f"붙임5 파일을 찾을 수 없습니다: {src}")

    return dest


def _merge_id_copies(price_folder, applicants, price):
    """붙임6: 배정자 실명확인증표 (신분증 합본, 앞면만)."""
    import database as db
    from processors.id_filter import filter_front_pages_only

    ap_ids = [ap['id'] for ap in applicants]
    id_docs = db.get_documents_for_applicant_ids(ap_ids, 'id_copy')

    if not id_docs:
        raise ValueError("신분증 파일이 없습니다")

    # 앞면만 필터링
    filtered_paths = []
    for doc in id_docs:
        fp = doc.get('file_path', '')
        if not os.path.isfile(fp):
            continue

        applicant_name = doc.get('name', '')
        filtered_path = filter_front_pages_only(fp, applicant_name)

        if filtered_path:
            filtered_paths.append(filtered_path)

    if not filtered_paths:
        raise ValueError("신분증 앞면 파일을 찾을 수 없습니다")

    output = os.path.join(price_folder, f'(붙임6) 배정자_실명확인증표_{price}원.pdf')
    merge_pdfs_in_order(filtered_paths, output)
    return output


def _merge_account_copies(price_folder, applicants, price):
    """붙임7: 배정자 증권계좌사본 합본."""
    import database as db

    ap_ids = [ap['id'] for ap in applicants]
    acct_docs = db.get_documents_for_applicant_ids(ap_ids, 'account_copy')

    if not acct_docs:
        raise ValueError("계좌사본 파일이 없습니다")

    file_paths = [d['file_path'] for d in acct_docs if os.path.isfile(d.get('file_path', ''))]
    if not file_paths:
        raise ValueError("계좌사본 파일을 찾을 수 없습니다")

    output = os.path.join(price_folder, f'(붙임7) 배정자_증권계좌사본_{price}원.pdf')
    merge_pdfs_in_order(file_paths, output)
    return output


def _copy_attachment9(price_folder):
    """붙임9: 법인등기부등본 복사."""
    src = os.path.join(TEMPLATES_DIR, 'reference_docs', '(붙임9) 법인등기부등본_에스투더블유 260408.pdf')
    dest = os.path.join(price_folder, '(붙임9) 법인등기부등본.pdf')
    shutil.copy2(src, dest)
    return dest
