# -*- coding: utf-8 -*-
"""
의무보유 Word 문서 자동 생성.
python-docx로 템플릿 표 행 교체 + 확약인 단락 교체.
"""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import os

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14_NS = 'http://schemas.microsoft.com/office/word/2010/wordml'


# ── 공통 헬퍼 ─────────────────────────────────────────────────────────────────

def _set_cell_text(cell, text: str):
    """셀 내 텍스트 교체. 첫 번째 paragraph의 첫 run 형식 유지."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)


def _clean_ids(tr_element):
    """deepcopy한 행에서 paraId, textId 등 고유 ID 속성 제거."""
    for el in tr_element.iter():
        el.attrib.pop(f'{{{W14_NS}}}paraId', None)
        el.attrib.pop(f'{{{W14_NS}}}textId', None)
        el.attrib.pop(qn('w:rsidR'), None)
        el.attrib.pop(qn('w:rsidRPr'), None)
        el.attrib.pop(qn('w:rsidTr'), None)
    return tr_element


def _clone_row(template_tr):
    new_tr = deepcopy(template_tr)
    _clean_ids(new_tr)
    return new_tr


def _tc_list(tr_element):
    """행의 실제 <w:tc> 요소 목록 반환 (gridSpan 포함 원본 요소)."""
    return tr_element.findall(f'{{{W_NS}}}tc')


def _set_vmerge(tc_element, mode):
    """
    tc 요소의 vMerge 속성 설정.
    mode: 'restart' | 'continue' | 'none'
    """
    tcPr = tc_element.find(f'{{{W_NS}}}tcPr')
    if tcPr is None:
        tcPr = etree.SubElement(tc_element, f'{{{W_NS}}}tcPr')
        tc_element.insert(0, tcPr)

    existing = tcPr.find(f'{{{W_NS}}}vMerge')
    if mode == 'none':
        if existing is not None:
            tcPr.remove(existing)
        return

    if existing is None:
        existing = etree.SubElement(tcPr, f'{{{W_NS}}}vMerge')

    if mode == 'restart':
        existing.set(f'{{{W_NS}}}val', 'restart')
    else:  # continue
        # vMerge without val attribute = continue
        existing.attrib.pop(f'{{{W_NS}}}val', None)


def _make_vmerge_continue(tc_element):
    """셀을 vMerge continue 상태로 변경하고 내용을 비움."""
    _set_vmerge(tc_element, 'continue')
    for p in tc_element.findall(f'{{{W_NS}}}p'):
        for child in list(p):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('r', 'proofErr', 'bookmarkStart', 'bookmarkEnd', 'hyperlink', 'del', 'ins'):
                p.remove(child)


def _merge_first_two_cells(tr_element):
    """합계 행의 첫 두 tc를 gridSpan=2로 병합하고 두 번째 tc 제거."""
    tcs = _tc_list(tr_element)
    if len(tcs) < 2:
        return
    tc0, tc1 = tcs[0], tcs[1]

    tcPr0 = tc0.find(f'{{{W_NS}}}tcPr')
    if tcPr0 is None:
        tcPr0 = etree.SubElement(tc0, f'{{{W_NS}}}tcPr')

    # gridSpan=2 설정
    gs = tcPr0.find(f'{{{W_NS}}}gridSpan')
    if gs is None:
        gs = etree.SubElement(tcPr0, f'{{{W_NS}}}gridSpan')
    gs.set(f'{{{W_NS}}}val', '2')

    # 너비 = tc0 + tc1 합산
    tcW0 = tcPr0.find(f'{{{W_NS}}}tcW')
    tcPr1 = tc1.find(f'{{{W_NS}}}tcPr')
    if tcW0 is not None and tcPr1 is not None:
        tcW1 = tcPr1.find(f'{{{W_NS}}}tcW')
        if tcW1 is not None:
            w0 = int(tcW0.get(f'{{{W_NS}}}w', 0) or 0)
            w1 = int(tcW1.get(f'{{{W_NS}}}w', 0) or 0)
            tcW0.set(f'{{{W_NS}}}w', str(w0 + w1))

    tr_element.remove(tc1)


def _replace_para_text(para, new_text: str):
    for i, run in enumerate(para.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ''


def _find_para_containing(doc, keyword: str):
    return [i for i, p in enumerate(doc.paragraphs) if keyword in p.text]


# ── 확약서 생성 ────────────────────────────────────────────────────────────────

def generate_hwakjakseo(template_path: str, output_path: str,
                        subjects: list, holding_start: str, holding_end: str):
    """
    의무보유 확약서 생성.

    subjects: list of dict {name, relationship, quantity, note}
    holding_start / holding_end: 'YYYY.MM.DD' 형식

    표 구조 (data rows):
      col0: 주주명      — 행별 개별
      col1: 관계        — 행별 개별
      col2: 종류        — 행별 개별 ('보통주')
      col3: 주식수      — 행별 개별
      col4: 의무보유사유 — 전체 vMerge, 첫 행만 '주1)'
      col5: 시작일      — 전체 vMerge, 첫 행만 날짜
      col6-8: 종료일(gridSpan=3) — 전체 vMerge, 첫 행만 날짜
      col9: 비고        — 전체 vMerge, 첫 행만 '-'
    합계 행:
      col0+col1 병합 → '합계', col2 → '보통주', col3 → 총주식수
    """
    REASON = '코스닥시장 상장규정 제26조 제1항 제6호 및 동항 단서조항에 의거하여 의무보유함'

    with open(template_path, 'rb') as f:
        doc = Document(f)

    # ── 표1 (index 1) 처리 ───────────────────────────────────────────────────
    table = doc.tables[1]
    rows = table.rows
    n = len(rows)
    DATA_START = 2

    if n <= DATA_START:
        raise ValueError(f"확약서 표 행이 부족합니다 (총 {n}행)")

    # 삭제 전 템플릿 저장
    template_data = deepcopy(rows[DATA_START]._tr)  # 첫 데이터 행 (row2)
    template_sum  = deepcopy(rows[n - 2]._tr)        # 합계 행 (row8)
    template_foot = deepcopy(rows[n - 1]._tr) if n >= DATA_START + 3 else None  # 빈 행 (row9)

    for row in list(rows[DATA_START:]):
        row._tr.getparent().remove(row._tr)

    total_qty = 0

    for i, subj in enumerate(subjects):
        qty = subj.get('quantity') or 0
        total_qty += qty
        is_first = (i == 0)

        new_tr = _clone_row(template_data)
        table._tbl.append(new_tr)
        new_row  = table.rows[-1]
        cells    = new_row.cells          # grid 기준 (cells[6][7][8] = 같은 tc)
        tcs      = _tc_list(new_tr)       # 실제 tc 요소 8개

        # 개별 열 (col0-3)
        _set_cell_text(cells[0], subj.get('name', ''))
        _set_cell_text(cells[1], subj.get('relationship', '미등기임원'))
        _set_cell_text(cells[2], '보통주')
        _set_cell_text(cells[3], f"{qty:,}")

        # 공통 열 (col4=사유, col5=시작일, tc[6]=종료일, tc[7]=비고)
        # 첫 행: vMerge restart + 내용 기입
        # 이후 행: vMerge continue + 내용 비움
        if is_first:
            _set_vmerge(tcs[4], 'restart')
            _set_cell_text(cells[4], '주1)')
            _set_vmerge(tcs[5], 'restart')
            _set_cell_text(cells[5], holding_start)
            _set_vmerge(tcs[6], 'restart')
            _set_cell_text(cells[6], holding_end)   # gridSpan=3이지만 첫 tc에 씀
            _set_vmerge(tcs[7], 'restart')
            _set_cell_text(cells[9], '-')            # cells[9] = tc[7] (비고)
        else:
            _make_vmerge_continue(tcs[4])
            _make_vmerge_continue(tcs[5])
            _make_vmerge_continue(tcs[6])
            _make_vmerge_continue(tcs[7])

    # ── 합계 행 ──────────────────────────────────────────────────────────────
    # col0+col1 병합 → '합계', col2 → '보통주', col3 → 총주식수
    sum_tr = _clone_row(template_sum)
    _merge_first_two_cells(sum_tr)          # tc[0] gridSpan=2, tc[1] 제거
    table._tbl.append(sum_tr)
    sum_row = table.rows[-1]
    sc = sum_row.cells                      # cells[0],[1] = tc[0]; cells[2] = 구 tc[2]
    _set_cell_text(sc[0], '합계')
    _set_cell_text(sc[2], '보통주')         # 종류 열
    _set_cell_text(sc[3], f'{total_qty:,}') # 주식수 열
    for ci in range(4, len(sc)):
        _set_cell_text(sc[ci], '')

    # 빈 footer 행 복원
    if template_foot is not None:
        table._tbl.append(_clone_row(template_foot))

    # ── 주1) 단락 업데이트 (표 바로 아래) ────────────────────────────────────
    for para in doc.paragraphs:
        if para.text.strip().startswith('주1)'):
            _replace_para_text(para, f'주1) {REASON}')
            break

    # ── 확약인 단락 교체 ──────────────────────────────────────────────────────
    name_para_indices = _find_para_containing(doc, '주주명 (본 인)')

    for i, subj in enumerate(subjects):
        name = subj.get('name', '')
        spaced = ' '.join(name) if all(ord(c) > 127 for c in name.replace(' ', '')) else name
        new_text = f'주주명 (본 인)   {spaced}   (인)'
        if i < len(name_para_indices):
            _replace_para_text(doc.paragraphs[name_para_indices[i]], new_text)

    # 남는 확약인 블록 제거
    if len(subjects) < len(name_para_indices):
        for idx in reversed(name_para_indices[len(subjects):]):
            for offset in range(4, -1, -1):
                target = idx - offset
                if 0 <= target < len(doc.paragraphs):
                    p = doc.paragraphs[target]
                    if p.text.strip() in ('확약인', '주식회사 에스투더블유', '') or \
                       '주주명 (본 인)' in p.text:
                        p._element.getparent().remove(p._element)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ── 계속보유신청 공문 생성 ────────────────────────────────────────────────────

def generate_gongmun(template_path: str, output_path: str,
                     date_str: str, processing_date: str,
                     subjects_dogok: list, subjects_yeouido: list,
                     holding_start: str, holding_end: str,
                     applicant_accounts: dict = None):
    """
    계속보유(의무보유)신청 공문 생성.

    date_str / processing_date: 'YYYY. MM. DD' 형식
    subjects_*: list of dict {name, account_number, quantity, branch, note}
    applicant_accounts: {name: [{account_number, quantity}, ...]} — 다중계좌 지원
    """
    with open(template_path, 'rb') as f:
        doc = Document(f)

    if applicant_accounts is None:
        applicant_accounts = {}

    # 날짜 단락 업데이트
    for para in doc.paragraphs:
        if '일    자' in para.text and ('2026' in para.text or '일    자 :' in para.text):
            _replace_para_text(para, f'일    자 : {date_str}')
        elif '예약처리일' in para.text:
            _replace_para_text(para, f'5. 예약처리일 : {processing_date}')

    REASON = '코스닥시장 상장규정 제26조 제1항 제6호 및 동항 단서조항에 의거하여 거래소와 협의'
    period_str = f'{holding_start}~{holding_end}'

    def _fill_table(table, subjects):
        """
        공문 표 채우기.

        열 구조:
          col0: 이름      (사람별 vMerge)
          col1: 계좌번호  (행별 개별)
          col2: 주식수    (행별 개별)
          col3: 기간      (전체 데이터 vMerge — 첫 행에 내용, 나머지 continue)
          col4: 사유      (전체 데이터 vMerge)
          col5: 지점      (전체 데이터 vMerge)
          col6: 비고      (전체 데이터 vMerge)
          col7: 여백
        """
        rows = table.rows
        n = len(rows)
        DATA_START = 2
        if n <= DATA_START:
            return

        # 삭제 전 행 템플릿 저장
        # row2: 첫 데이터행 (col0=restart, col3-6=restart)
        # row3: 같은 사람 두번째 계좌 (col0=continue, col3-6=continue)
        # row4: 새 사람 단일 계좌 (col0=none, col3-6=continue)
        # row9(마지막): 합계행 (gridSpan 다름)
        tpl_first      = deepcopy(rows[DATA_START]._tr)     # row2
        tpl_same_cont  = deepcopy(rows[DATA_START + 1]._tr) if n > DATA_START + 1 else deepcopy(rows[DATA_START]._tr)  # row3
        tpl_new_person = deepcopy(rows[DATA_START + 2]._tr) if n > DATA_START + 2 else deepcopy(rows[DATA_START]._tr)  # row4
        tpl_sum        = deepcopy(rows[-1]._tr)              # 합계행

        # 전부 제거
        for row in list(rows[DATA_START:]):
            row._tr.getparent().remove(row._tr)

        total_qty = 0
        is_first_row = True  # 전체 데이터에서 첫 행 여부 (기간/사유/지점/비고 restart 결정)

        for subj in subjects:
            name = subj.get('name', '')
            branch = subj.get('branch', '도곡')
            branch_name = '도곡WM센터' if branch == '도곡' else '여의도금융1센터'
            note = subj.get('note', '주1)') or '주1)'

            # 해당 사람의 계좌 목록: applicant_accounts에서 가져오거나 없으면 holding_subjects 데이터 사용
            acc_list = applicant_accounts.get(name)
            if not acc_list:
                acc_list = [{
                    'account_number': subj.get('account_number', ''),
                    'quantity': subj.get('quantity') or 0,
                }]

            n_acc = len(acc_list)

            for acc_idx, acc in enumerate(acc_list):
                qty = acc.get('quantity') or 0
                account = acc.get('account_number', '')
                total_qty += qty

                # 어떤 행 템플릿을 쓸지 결정
                if is_first_row:
                    # 첫 데이터행: col0=restart, col3-6=restart (tpl_first 그대로)
                    new_tr = _clone_row(tpl_first)
                    # 만약 단일 계좌인 경우 col0 vMerge=restart지만 continue가 없으므로 none으로 변경
                    if n_acc == 1:
                        tcs = _tc_list(new_tr)
                        if tcs:
                            _set_vmerge(tcs[0], 'none')
                    is_first_row = False
                elif acc_idx == 0:
                    if n_acc > 1:
                        # 새 사람 (첫 번째 아님), 다중 계좌: col0=restart, col3-6=continue
                        new_tr = _clone_row(tpl_new_person)
                        tcs = _tc_list(new_tr)
                        if tcs:
                            _set_vmerge(tcs[0], 'restart')
                    else:
                        # 새 사람, 단일 계좌: col0=none, col3-6=continue
                        new_tr = _clone_row(tpl_new_person)
                else:
                    # 같은 사람 추가 계좌: col0=continue, col3-6=continue
                    new_tr = _clone_row(tpl_same_cont)

                table._tbl.append(new_tr)
                new_row = table.rows[-1]
                cells = new_row.cells

                # 이름: continue 행이면 건드리지 않음 (empty vMerge 셀)
                tcs = _tc_list(new_tr)
                col0_vm = tcs[0].find(f'{{{W_NS}}}tcPr/{{{W_NS}}}vMerge') if tcs else None
                is_name_visible = (col0_vm is None) or (col0_vm.get(f'{{{W_NS}}}val') == 'restart')

                if is_name_visible and len(cells) > 0:
                    _set_cell_text(cells[0], name)

                if len(cells) > 1: _set_cell_text(cells[1], account)
                if len(cells) > 2: _set_cell_text(cells[2], f'{qty:,}')

                # 기간/사유/지점/비고: restart 행에만 내용 씀
                col3_vm = tcs[3].find(f'{{{W_NS}}}tcPr/{{{W_NS}}}vMerge') if len(tcs) > 3 else None
                is_period_visible = (col3_vm is None) or (col3_vm.get(f'{{{W_NS}}}val') == 'restart')

                if is_period_visible:
                    if len(cells) > 3: _set_cell_text(cells[3], period_str)
                    if len(cells) > 4: _set_cell_text(cells[4], REASON)
                    if len(cells) > 5: _set_cell_text(cells[5], branch_name)
                    if len(cells) > 6: _set_cell_text(cells[6], note)

        # 합계 행 (원본 구조: gridSpan=2 '합계' | qty | gridSpan=4 '' | '')
        sum_tr = _clone_row(tpl_sum)
        table._tbl.append(sum_tr)
        sum_row = table.rows[-1]
        sum_tcs = _tc_list(sum_tr)
        if sum_tcs:
            # tc[0] = '합계' (gridSpan=2)
            _set_cell_text(sum_row.cells[0], '합계')
            # tc[1] = 총주식수 (gridSpan=1)
            if len(sum_tcs) > 1:
                _set_cell_text(sum_row.cells[2], f'{total_qty:,}')  # cells[2] = gridSpan=2인 tc[0] 다음 첫 단독셀
            # 나머지 빈칸
            for ci in range(len(sum_row.cells)):
                if ci not in (0, 1, 2):
                    _set_cell_text(sum_row.cells[ci], '')

    # 공문: doc.tables[0]=도곡, doc.tables[1]=여의도
    if len(doc.tables) >= 1:
        all_subjects = subjects_dogok + subjects_yeouido
        _fill_table(doc.tables[0], subjects_dogok if subjects_dogok else all_subjects)
    if len(doc.tables) >= 2:
        _fill_table(doc.tables[1], subjects_yeouido if subjects_yeouido else [])

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
